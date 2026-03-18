"""Create a sample DOCX resume for smoke testing."""

from __future__ import annotations

import json
from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from build_log_utils import log_file_event
from docx import Document
from docx.shared import Inches, Pt


def build_sample_resume(schema_path: str | Path, output_path: str | Path) -> str:
    """Generate a realistic sample resume DOCX from the JSON schema."""
    data = json.loads(Path(schema_path).read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.19)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    name_p = document.add_paragraph()
    name_run = name_p.add_run(data["name"])
    name_run.bold = True
    name_run.font.size = Pt(14)

    summary_p = document.add_paragraph()
    summary_run = summary_p.add_run(data["summary"])
    summary_run.font.size = Pt(10)

    skills_heading = document.add_paragraph()
    skills_heading.add_run("Technical Skills").bold = True

    for category, values in data["skills"].items():
        skills_p = document.add_paragraph()
        label = skills_p.add_run(f"{category}: ")
        label.bold = True
        label.font.size = Pt(10)
        values_run = skills_p.add_run(", ".join(values))
        values_run.font.size = Pt(10)

    exp_heading = document.add_paragraph()
    exp_heading.add_run("Experience").bold = True

    for exp in data["experience"]:
        header_p = document.add_paragraph()
        header_run = header_p.add_run(f'{exp["role"]} | {exp["company"]}\t{exp["dates"]}')
        header_run.bold = True
        header_run.font.size = Pt(10)
        for idx, bullet in enumerate(exp["bullets"]):
            bullet_p = document.add_paragraph(style="List Bullet")
            bullet_run = bullet_p.add_run(bullet)
            bullet_run.font.size = Pt(10)
            if idx == 0:
                continuation_p = document.add_paragraph()
                continuation_run = continuation_p.add_run(
                    "Led collaboration with analysts and engineering partners to operationalize the solution."
                )
                continuation_run.font.size = Pt(10)

    edu_heading = document.add_paragraph()
    edu_heading.add_run("Education").bold = True
    for edu in data["education"]:
        p = document.add_paragraph()
        run = p.add_run(f'{edu["degree"]} | {edu["institution"]}\t{edu["dates"]}')
        run.font.size = Pt(10)

    proj_heading = document.add_paragraph()
    proj_heading.add_run("Projects").bold = True
    for project in data["projects"]:
        name_p = document.add_paragraph()
        name_run = name_p.add_run(project["name"])
        name_run.bold = True
        name_run.font.size = Pt(10)
        for bullet in project["bullets"]:
            p = document.add_paragraph(style="List Bullet")
            run = p.add_run(bullet)
            run.font.size = Pt(10)

    cert_heading = document.add_paragraph()
    cert_heading.add_run("Certifications").bold = True
    cert_p = document.add_paragraph()
    cert_run = cert_p.add_run(" | ".join(data["certifications"]))
    cert_run.font.size = Pt(10)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    existed = output_path.exists()
    document.save(output_path)
    log_file_event(
        "MODIFIED" if existed else "CREATED",
        output_path,
        "Generated sample resume DOCX for smoke testing",
    )
    return str(output_path)


if __name__ == "__main__":
    target = Path("data/sample_resume.docx")
    print(build_sample_resume("data/resume_schema.json", target))
