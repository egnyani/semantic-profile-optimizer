"""Create the base resume template DOCX."""

from __future__ import annotations

from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from build_log_utils import log_file_event
from docx import Document
from docx.shared import Inches, Pt


def create_template(output_path: str | Path) -> str:
    """Generate a base DOCX template with the required page layout."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.19)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    paragraph = document.add_paragraph()
    run = paragraph.add_run("Resume Template")
    run.bold = True
    run.font.size = Pt(10)

    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    existed = target.exists()
    document.save(target)
    log_file_event(
        "MODIFIED" if existed else "CREATED",
        target,
        "Generated base DOCX template with required margins",
    )
    return str(target)


if __name__ == "__main__":
    print(create_template("templates/resume_template.docx"))
