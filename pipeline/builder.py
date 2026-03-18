"""DOCX builder for tailored resumes."""

from __future__ import annotations

import math
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


def _sanitize_filename_part(text: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return cleaned or "Unknown"


def make_output_filename(company: str, role: str) -> str:
    """Build the required output filename pattern."""
    return f"Gnyani_{_sanitize_filename_part(company)}_{_sanitize_filename_part(role)}.docx"


def estimate_line_count(data: dict[str, Any]) -> int:
    """Estimate rendered line count for one-page enforcement."""
    lines = 0
    if data.get("name"):
        lines += 2
    if data.get("summary"):
        lines += math.ceil(len(data["summary"]) / 85)
    if data.get("skills"):
        for key, values in data["skills"].items():
            skill_text = f"{key}: {', '.join(values)}"
            lines += max(1, math.ceil(len(skill_text) / 95))
    for exp in data.get("experience", []):
        lines += 2 if exp.get("company") else 1
        for bullet in exp.get("bullets", []):
            lines += max(1, math.ceil(len(bullet) / 80))
    lines += len(data.get("education", []))
    for project in data.get("projects", []):
        lines += 1
        for bullet in project.get("bullets", []):
            lines += max(1, math.ceil(len(bullet) / 80))
    lines += len(data.get("certifications", []))
    lines += 8
    return lines


def apply_one_page_trimming(data: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
    """Trim least-relevant bullets until the document likely fits on one page."""
    trimmed = deepcopy(data)
    warnings: list[str] = []
    experiences = trimmed.get("experience", [])
    if not experiences:
        return trimmed, warnings

    protected_index = 0
    order: list[int] = []
    for exp_index in range(len(experiences) - 1, -1, -1):
        if exp_index == protected_index:
            continue
        company = str(experiences[exp_index].get("company", "")).lower()
        role = str(experiences[exp_index].get("role", "")).lower()
        if "acmesia" in company and exp_index not in order:
            order.append(exp_index)
        if "intern" in role and exp_index not in order:
            order.append(exp_index)
    for exp_index in range(len(experiences) - 1, -1, -1):
        if exp_index != protected_index and exp_index not in order:
            order.append(exp_index)

    while estimate_line_count(trimmed) > 60:
        removed = False
        for exp_index in order:
            exp = trimmed["experience"][exp_index]
            if len(exp.get("bullets", [])) > 1:
                removed_bullet = exp["bullets"].pop()
                warnings.append(
                    f'Trimmed bullet from {exp.get("company", "experience")}: {removed_bullet}'
                )
                removed = True
                break
        if not removed:
            break
    return trimmed, warnings


def _set_run_font(run, bold: bool = False) -> None:
    run.font.size = Pt(10)
    run.bold = bold


def _set_paragraph_spacing(paragraph, before: float = 0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(0)


def add_bottom_border(paragraph) -> None:
    """Add a thin bottom border to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    for child in p_pr.findall(qn("w:pBdr")):
        p_pr.remove(child)
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    _set_paragraph_spacing(paragraph, before=4)
    run = paragraph.add_run(title)
    _set_run_font(run, bold=True)
    add_bottom_border(paragraph)


def _add_bullet_paragraph(document: Document, text: str, is_first: bool) -> None:
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.paragraph_format.left_indent = Inches(0.25)
    _set_paragraph_spacing(paragraph, before=2 if is_first else 0.1)
    run = paragraph.add_run(text)
    _set_run_font(run)


def _load_base_document(template_path: str | Path | None) -> Document:
    if template_path and Path(template_path).exists():
        return Document(str(template_path))
    return Document()


def json_to_docx(
    data: dict[str, Any],
    output_path: str | Path,
    template_path: str | Path | None = None,
) -> tuple[str, list[str]]:
    """Rebuild a resume DOCX from structured data."""
    trimmed_data, warnings = apply_one_page_trimming(data)
    document = _load_base_document(template_path)

    if document.sections:
        section = document.sections[0]
    else:
        section = document.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.19)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    while document.paragraphs:
        p = document.paragraphs[0]._element
        p.getparent().remove(p)

    name_p = document.add_paragraph()
    name_p.style = document.styles["Title"]
    _set_paragraph_spacing(name_p)
    name_run = name_p.add_run(trimmed_data.get("name", "Candidate"))

    if trimmed_data.get("summary"):
        summary_p = document.add_paragraph()
        summary_p.style = document.styles["Normal"]
        _set_paragraph_spacing(summary_p)
        summary_run = summary_p.add_run(trimmed_data["summary"])
        _set_run_font(summary_run)

    if trimmed_data.get("experience"):
        _add_section_heading(document, "Experience")
        for exp in trimmed_data["experience"]:
            company_line = exp.get("company", "")
            location = exp.get("location", "")
            if company_line:
                company_p = document.add_paragraph()
                company_p.style = document.styles["Normal"]
                _set_paragraph_spacing(company_p, before=4)
                company_text = f"{company_line}\t{location}".strip()
                company_run = company_p.add_run(company_text)
                _set_run_font(company_run, bold=True)

            role_p = document.add_paragraph()
            role_p.style = document.styles["Normal"]
            _set_paragraph_spacing(role_p, before=1)
            role_text = f'{exp.get("role", "")}\t{exp.get("dates", "")}'.strip()
            role_run = role_p.add_run(role_text)
            _set_run_font(role_run, bold=True)

            for bullet_index, bullet in enumerate(exp.get("bullets", [])):
                _add_bullet_paragraph(document, bullet, is_first=bullet_index == 0)

    if trimmed_data.get("education"):
        _add_section_heading(document, "Education")
        for edu in trimmed_data["education"]:
            p = document.add_paragraph()
            p.style = document.styles["Normal"]
            _set_paragraph_spacing(p, before=1)
            left_text = edu.get("degree", "")
            if edu.get("institution"):
                left_text = f'{left_text}, {edu.get("institution", "")}'
            text = f'{left_text}\t{edu.get("dates", "")}'
            run = p.add_run(text.strip())
            _set_run_font(run, bold=True)

    if trimmed_data.get("skills"):
        _add_section_heading(document, "Technical Skills")
        for category, values in trimmed_data["skills"].items():
            p = document.add_paragraph()
            p.style = document.styles["Normal"]
            _set_paragraph_spacing(p)
            label_run = p.add_run(f"{category}: ")
            _set_run_font(label_run, bold=True)
            values_run = p.add_run(", ".join(values))
            _set_run_font(values_run)

    if trimmed_data.get("projects"):
        _add_section_heading(document, "Projects")
        for project in trimmed_data["projects"]:
            name_p = document.add_paragraph()
            name_p.style = document.styles["Normal"]
            _set_paragraph_spacing(name_p, before=1)
            name_run = name_p.add_run(project.get("name", "Project"))
            _set_run_font(name_run, bold=True)
            for bullet_index, bullet in enumerate(project.get("bullets", [])):
                _add_bullet_paragraph(document, bullet, is_first=bullet_index == 0)

    if trimmed_data.get("certifications"):
        _add_section_heading(document, "Certifications")
        cert_p = document.add_paragraph()
        cert_p.style = document.styles["Normal"]
        _set_paragraph_spacing(cert_p)
        cert_run = cert_p.add_run(" | ".join(trimmed_data["certifications"]))
        _set_run_font(cert_run)

    output_path = str(output_path)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path, warnings
